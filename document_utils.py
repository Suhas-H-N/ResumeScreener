"""
Document processing utilities
Enhanced PDF and DOCX text extraction
"""
import os
import logging
from pathlib import Path
from typing import Optional
import PyPDF2
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handle document text extraction"""
    
    ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'doc'}
    
    @staticmethod
    def allowed_file(filename: str) -> bool:
        """Check if file extension is allowed"""
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in DocumentProcessor.ALLOWED_EXTENSIONS
    
    @staticmethod
    def extract_text(filepath: str) -> str:
        """
        Extract text from document
        
        Args:
            filepath: Path to the document
            
        Returns:
            Extracted text content
        """
        if not os.path.exists(filepath):
            logger.error(f"File not found: {filepath}")
            return ""
        
        ext = Path(filepath).suffix.lower()
        
        try:
            if ext == '.pdf':
                return DocumentProcessor._extract_from_pdf(filepath)
            elif ext == '.docx':
                return DocumentProcessor._extract_from_docx(filepath)
            elif ext in ['.txt', '.doc']:
                return DocumentProcessor._extract_from_txt(filepath)
            else:
                logger.warning(f"Unsupported file type: {ext}")
                return ""
        except Exception as e:
            logger.error(f"Text extraction error for {filepath}: {e}")
            return ""
    
    @staticmethod
    def _extract_from_pdf(filepath: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                # Extract from all pages
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text() or ""
                        text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num}: {e}")
                        continue
                
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
        
        return text.strip()
    
    @staticmethod
    def _extract_from_docx(filepath: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(filepath)
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(' '.join(row_text))
            
            # Combine all text
            all_text = paragraphs + table_text
            return '\n'.join(all_text)
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ""
    
    @staticmethod
    def _extract_from_txt(filepath: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            return ""
    
    @staticmethod
    def validate_resume_content(text: str, min_length: int = 50) -> tuple[bool, str]:
        """
        Validate if the extracted text is a valid resume
        
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not text or len(text.strip()) == 0:
            return False, "No text could be extracted from the document"
        
        if len(text) < min_length:
            return False, f"Resume text is too short ({len(text)} characters). Please upload a complete resume."
        
        # Check for common resume indicators
        text_lower = text.lower()
        resume_indicators = [
            'experience', 'education', 'skills', 'work', 'employment',
            'university', 'college', 'bachelor', 'master', 'degree',
            'project', 'achievement', 'certification'
        ]
        
        found_indicators = sum(1 for indicator in resume_indicators if indicator in text_lower)
        
        if found_indicators < 2:
            return False, "This doesn't appear to be a resume. Please upload a valid resume document."
        
        return True, ""
    
    @staticmethod
    def extract_contact_info(text: str) -> dict:
        """
        Extract contact information from resume
        
        Returns:
            Dictionary with email, phone, linkedin, etc.
        """
        import re
        
        contact = {
            'email': None,
            'phone': None,
            'linkedin': None,
            'github': None,
        }
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact['email'] = emails[0]
        
        # Extract phone
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)
        if phones:
            contact['phone'] = ''.join(phones[0]) if isinstance(phones[0], tuple) else phones[0]
        
        # Extract LinkedIn
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_matches = re.findall(linkedin_pattern, text.lower())
        if linkedin_matches:
            contact['linkedin'] = f"https://{linkedin_matches[0]}"
        
        # Extract GitHub
        github_pattern = r'github\.com/[\w-]+'
        github_matches = re.findall(github_pattern, text.lower())
        if github_matches:
            contact['github'] = f"https://{github_matches[0]}"
        
        return contact
    
    @staticmethod
    def extract_education(text: str) -> list:
        """Extract education information"""
        import re
        
        degrees = []
        degree_pattern = r'(bachelor|master|phd|doctorate|associate|b\.s\.|m\.s\.|b\.a\.|m\.a\.|mba).*?(?:degree|of|in)\s+([\w\s]+)'
        
        matches = re.findall(degree_pattern, text.lower())
        for match in matches:
            degrees.append(f"{match[0].title()} in {match[1].strip().title()}")
        
        return degrees[:5]  # Return up to 5 degrees
